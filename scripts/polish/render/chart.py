"""Chart renderer — matplotlib → PNG → embedded inline image.

Background: the python-docx fork we run against does not ship `docx.chart`,
so we can't emit native Word charts from Python. Instead, we render charts
with matplotlib using the brand palette and embed the result as a centered
inline PNG. The data lives on the canonical Chart block, so the output is
still deterministic and rebuilt from scratch every run — it just isn't an
editable Word chart inside the resulting .docx.

This is a pragmatic tradeoff vs. the alternative of hand-authoring DrawingML
chart XML (brittle, complex). Visual parity with the Figma design is the
goal, and a clean matplotlib render in brand colors meets it.
"""
from __future__ import annotations

from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from ..model import Chart, Series
from . import tokens as T
from .xml_utils import apply_text_style, set_paragraph_spacing


def render(doc_docx, chart: Chart) -> None:
    png = _render_png(chart)
    if png is None:
        _render_unavailable(doc_docx, chart, reason="matplotlib render failed")
        return

    # Title (above image, branded)
    if chart.title:
        tp = doc_docx.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(tp, before_twips=120, after_twips=60, line_multiple=1.25)
        trun = tp.add_run(chart.title)
        apply_text_style(trun, T.TITLES_SUB)

    # Centered image
    p = doc_docx.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_twips=60, after_twips=120, line_multiple=1.0)
    run = p.add_run()
    run.add_picture(BytesIO(png), width=Inches(6.2))


def _render_png(chart: Chart) -> bytes | None:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return None

    try:
        fig, ax = plt.subplots(figsize=(9.5, 4.8), dpi=160)
        _draw(ax, chart)
        _brand_axes(ax, chart)
        buf = BytesIO()
        fig.tight_layout()
        fig.savefig(buf, format="png", bbox_inches="tight",
                    facecolor="white", edgecolor="none")
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
        return None


# ── drawing ─────────────────────────────────────────────────────────────────

def _palette_hex() -> list[str]:
    return [f"#{T.hex_color(c)}" for c in T.CHART_SERIES_COLORS]


def _draw(ax, chart: Chart) -> None:
    colors = _palette_hex()
    kind = chart.kind
    cats = list(chart.categories)
    series = list(chart.series)

    if kind == "line":
        for i, s in enumerate(series):
            ax.plot(cats, _pad(s.values, len(cats)),
                    marker="o", color=colors[i % len(colors)],
                    linewidth=2.2, markersize=6, label=s.name or "")
        ax.grid(True, axis="y", color="#E5E7E9", linewidth=0.8)
        return

    if kind in ("pie", "donut"):
        vals = _pad(series[0].values, len(cats)) if series else []
        if not vals:
            return
        wedgeprops = {"width": 0.45} if kind == "donut" else {}
        ax.pie(vals, labels=cats,
               colors=colors[: len(cats)],
               autopct="%1.0f%%",
               wedgeprops=wedgeprops,
               textprops={"color": "#202124", "fontsize": 10})
        ax.axis("equal")
        return

    if kind == "funnel":
        s = series[0] if series else Series(name="", values=[])
        vals = _pad(s.values, len(cats))
        pairs = sorted(zip(cats, vals), key=lambda p: -p[1])
        cats = [p[0] for p in pairs]
        vals = [p[1] for p in pairs]
        ax.barh(cats, vals, color=colors[0])
        ax.invert_yaxis()
        ax.grid(True, axis="x", color="#E5E7E9", linewidth=0.8)
        return

    if kind == "stacked":
        import numpy as np
        x = np.arange(len(cats))
        bottom = np.zeros(len(cats))
        for i, s in enumerate(series):
            vals = _pad(s.values, len(cats))
            ax.bar(x, vals, bottom=bottom,
                   color=colors[i % len(colors)], label=s.name or "")
            bottom = bottom + np.array(vals, dtype=float)
        ax.set_xticks(x)
        ax.set_xticklabels(cats)
        ax.grid(True, axis="y", color="#E5E7E9", linewidth=0.8)
        return

    if kind == "bar":
        # Horizontal single or grouped
        _draw_grouped(ax, cats, series, colors, horizontal=True)
        ax.grid(True, axis="x", color="#E5E7E9", linewidth=0.8)
        return

    # column / other → vertical grouped bars
    _draw_grouped(ax, cats, series, colors, horizontal=False)
    ax.grid(True, axis="y", color="#E5E7E9", linewidth=0.8)


def _draw_grouped(ax, cats, series, colors, *, horizontal: bool) -> None:
    import numpy as np
    n = len(cats)
    k = max(1, len(series))
    width = 0.8 / k
    x = np.arange(n)
    for i, s in enumerate(series):
        vals = _pad(s.values, n)
        offset = (i - (k - 1) / 2) * width
        positions = x + offset
        color = colors[i % len(colors)]
        if horizontal:
            ax.barh(positions, vals, height=width, color=color,
                    label=s.name or "")
        else:
            ax.bar(positions, vals, width=width, color=color,
                   label=s.name or "")
    if horizontal:
        ax.set_yticks(x)
        ax.set_yticklabels(cats)
        ax.invert_yaxis()
    else:
        ax.set_xticks(x)
        ax.set_xticklabels(cats, rotation=0 if max(len(c) for c in cats) <= 6 else 25,
                           ha="center" if max(len(c) for c in cats) <= 6 else "right")


def _pad(values, n: int) -> list[float]:
    out = [float(v) if v is not None else 0.0 for v in values]
    if len(out) < n:
        out = out + [0.0] * (n - len(out))
    return out[:n]


def _brand_axes(ax, chart: Chart) -> None:
    if chart.kind in ("pie", "donut"):
        return
    for side in ("top", "right"):
        ax.spines[side].set_visible(False)
    for side in ("left", "bottom"):
        ax.spines[side].set_color("#D7DBDD")
    ax.tick_params(colors="#626567", labelsize=9)
    ax.set_axisbelow(True)
    if any(getattr(s, "name", "") for s in chart.series) and len(chart.series) > 1:
        ax.legend(
            loc="upper center",
            bbox_to_anchor=(0.5, -0.18),
            ncol=min(len(chart.series), 4),
            frameon=False,
            fontsize=9,
        )


def _render_unavailable(doc_docx, chart: Chart, reason: str) -> None:
    p = doc_docx.add_paragraph()
    set_paragraph_spacing(p, before_twips=120, after_twips=60, line_multiple=1.3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Chart · {chart.kind}] {(chart.title or '').strip()}")
    apply_text_style(run, T.TEXT_MAIN, override_bold=True)
    if reason:
        q = doc_docx.add_paragraph()
        set_paragraph_spacing(q, after_twips=120, line_multiple=1.2)
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q_run = q.add_run(f"Chart data unavailable — {reason}")
        apply_text_style(q_run, T.TEXT_DISCLAIMER, override_italic=True)
