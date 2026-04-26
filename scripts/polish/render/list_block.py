"""List renderer — bulleted and numbered lists with clean hanging indent.

Uses a proper hanging paragraph + explicit tab stop so the glyph and body
text align naturally regardless of Word's default tab stop width.

Indent geometry (level 0):
  - Left edge of body text: 400 twips (~0.28 in)
  - Glyph hangs back to: 120 twips (~0.08 in)
  - Tab stop at 400 twips snaps the body flush

Each nesting level adds 360 twips to the left/tab stop values.
"""
from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model import List, ListItem
from . import tokens as T
from .xml_utils import apply_text_style, set_paragraph_spacing


_BULLET_GLYPHS = ["•", "◦", "▸"]

_BASE_LEFT   = 200   # twips — where body text starts (level 0, ~0.14 in)
_HANGING     = 200   # twips — glyph at left margin (0 indent), text starts at _BASE_LEFT
_LEVEL_STEP  = 240   # twips added per nesting level


def render(doc_docx, list_block: List) -> None:
    counters: dict[int, int] = {}
    for item in list_block.items:
        _render_item(doc_docx, item, counters)


def _render_item(doc_docx, item: ListItem, counters: dict[int, int]) -> None:
    para = doc_docx.add_paragraph()
    level = max(0, min(item.level, 2))
    left = _BASE_LEFT + level * _LEVEL_STEP

    set_paragraph_spacing(para, before_twips=0, after_twips=80,
                          line_multiple=T.TEXT_MAIN.line_spacing)
    _set_hanging_indent(para, left_twips=left, hanging_twips=_HANGING, tab_stop_twips=left)

    if item.ordered:
        counters[level] = counters.get(level, 0) + 1
        for deeper in list(counters):
            if deeper > level:
                counters.pop(deeper, None)
        glyph = f"{counters[level]}."
    else:
        glyph = _BULLET_GLYPHS[level % len(_BULLET_GLYPHS)]

    glyph_run = para.add_run(f"{glyph}\t")
    apply_text_style(glyph_run, T.TEXT_MAIN)

    for r in item.runs:
        if not r.text:
            continue
        run = para.add_run(r.text)
        apply_text_style(run, T.TEXT_MAIN, override_bold=r.bold, override_italic=r.italic)


def _set_hanging_indent(para, *, left_twips: int, hanging_twips: int,
                        tab_stop_twips: int) -> None:
    pPr = para._p.get_or_add_pPr()

    # Indentation
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    pPr.append(ind)

    # Tab stop so \t snaps body text flush with left_twips
    for old in pPr.findall(qn("w:tabs")):
        pPr.remove(old)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(tab_stop_twips))
    tabs.append(tab)
    pPr.append(tabs)
