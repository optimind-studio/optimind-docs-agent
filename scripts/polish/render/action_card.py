"""Action-card renderer — numbered recommendation card with red left border.

Implementation note: title + body are placed inside a single borderless 1-cell
table whose cell carries the red left border, so the border is continuous
across the title/body break instead of fragmenting the way separate paragraph
borders do.
"""
from __future__ import annotations

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model import ActionCard
from . import tokens as T
from .xml_utils import (
    apply_text_style, set_cell_borders, set_cell_color,
    set_cell_padding, set_paragraph_spacing,
)


def render(doc_docx, card: ActionCard) -> None:
    # Spacer above the card.
    spacer = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer, before_twips=0, after_twips=80, line_multiple=1.0)

    # Single-cell wrapper table — gives the card one continuous left border.
    tbl = doc_docx.add_table(rows=1, cols=1)
    tbl.autofit = False
    _set_table_frame(tbl)

    cell = tbl.rows[0].cells[0]
    set_cell_color(cell, T.WHITE)
    # Borders: only the brand-red left border. ~2.25pt (18 eighths-of-a-point).
    set_cell_borders(cell, left=T.RED, size=18,
                     top=None, bottom=None, right=None)
    set_cell_padding(cell, top=120, bottom=160, left=200, right=80)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Drop the default empty paragraph so we can add ours cleanly.
    for p in list(cell.paragraphs):
        cell._tc.remove(p._p)

    # Title paragraph: "N. Bold title text"
    title_para = cell.add_paragraph()
    set_paragraph_spacing(title_para, before_twips=0, after_twips=40, line_multiple=1.3)

    num_run = title_para.add_run(f"{card.number}. ")
    apply_text_style(num_run, T.LABELS_MAIN, override_color=T.RED)

    title_run = title_para.add_run(card.title)
    apply_text_style(title_run, T.TEXT_MAIN, override_bold=True)

    # Body paragraph (inside the same cell — same continuous left border).
    if card.body:
        body_para = cell.add_paragraph()
        set_paragraph_spacing(body_para, before_twips=0, after_twips=0, line_multiple=1.3)
        body_run = body_para.add_run(card.body)
        apply_text_style(body_run, T.TEXT_MAIN, override_color=T.TEXT_SEC)


def _set_table_frame(tbl) -> None:
    """Borderless table frame at full page width."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for tag in ("w:tblStyle", "w:tblBorders", "w:tblCellMar"):
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
