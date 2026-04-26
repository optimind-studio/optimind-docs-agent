"""Table renderer — Classic (red header) + Minimal (rule-based) variants.

Builds the python-docx table from scratch using the canonical Table model.
Handles:
  - Multi-row headers (repeated for each row in table.headers)
  - Merged cells via MergeSpec (rowspan + colspan)
  - Numeric-column right alignment (auto-detected per column)
  - Zebra stripes in Classic; horizontal rules in Minimal

Port of pattern in scripts/polish_doc.py:431-598 — but operates on our typed model.
"""
from __future__ import annotations

import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..model import MergeSpec, Table
from . import tokens as T
from .xml_utils import (
    apply_text_style, set_cell_borders, set_cell_color, set_cell_grid_span,
    set_cell_padding, set_cell_vertical_merge, set_cell_width,
    set_paragraph_spacing,
)


_NUMERIC_RE = re.compile(r'^\s*[\$€£¥]?\s*[-+]?\d[\d,]*(\.\d+)?\s*%?\s*$')

_BADGE_POSITIVE = frozenset({
    "strong", "top performer", "good", "excellent ctr", "best in program",
    "scale up", "top opener", "lead2sale", "high", "excellent",
})
_BADGE_NEGATIVE = frozenset({
    "underperforming", "low open rate", "low", "poor", "weak",
})
_BADGE_NEUTRAL = frozenset({
    "average", "solid", "opens↑ ctr↓", "core market", "niche re-activation",
    "increase volume", "moderate", "mixed",
})
_BADGE_ALL = _BADGE_POSITIVE | _BADGE_NEGATIVE | _BADGE_NEUTRAL


def render(doc_docx, table: Table) -> None:
    if not table.rows and not table.headers:
        return

    # Spacer paragraph above the table so it breathes with surrounding content.
    spacer = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer, before_twips=0, after_twips=80, line_multiple=1.0)

    # Multi-row header fixup: some sources produce a merged spanning cell in
    # row 0 (a table title) with the column labels in row 1.  Ingest places
    # row 0 in headers and row 1 as the first body row, leaving rows[0] as
    # the label row.  Detect this: if rows[0] has no empty cells and looks
    # like labels (no numeric values) while headers[-1] is all-empty or a
    # single-cell merged span, promote rows[0] into headers.
    body_rows = list(table.rows)
    extra_headers: list = []
    if (table.headers and body_rows
            and _row_is_label_candidate(body_rows[0])
            and _row_is_spanning_title(table.headers[-1])):
        extra_headers = [body_rows.pop(0)]

    headers = list(table.headers) + extra_headers

    # Build a single matrix of all rows including header rows for rendering.
    header_row_count = len(headers)
    total_rows = header_row_count + len(body_rows)
    n_cols = _infer_col_count(table)

    if total_rows == 0 or n_cols == 0:
        return

    t = doc_docx.add_table(rows=total_rows, cols=n_cols)
    t.autofit = False
    _set_table_frame(t)

    # Populate cells
    for row_idx in range(total_rows):
        if row_idx < header_row_count:
            source_row = headers[row_idx]
            is_header = True
        else:
            source_row = body_rows[row_idx - header_row_count]
            is_header = False
        _populate_row(t, row_idx, source_row, n_cols,
                      is_header=is_header,
                      is_last_body=row_idx == total_rows - 1 and not is_header,
                      numeric_cols=_numeric_columns(body_rows, n_cols))

    # Apply merges. MergeSpec indexes into total_rows (0-based, headers first).
    for m in table.merges:
        _apply_merge(t, m, n_cols=n_cols, total_rows=total_rows)

    # Caption / trailing spacer
    if table.caption:
        _add_caption(doc_docx, table.caption)
    else:
        # Add a small spacer below the table even without a caption.
        post = doc_docx.add_paragraph()
        set_paragraph_spacing(post, before_twips=0, after_twips=160, line_multiple=1.0)


# ── multi-row header helpers ────────────────────────────────────────────────

def _row_is_spanning_title(row) -> bool:
    """True if a header row looks like a merged spanning title: exactly one
    non-empty cell (the rest are empty / continuation stubs)."""
    non_empty = [c for c in row if str(c).strip()]
    return len(non_empty) <= 1


def _row_is_label_candidate(row) -> bool:
    """True if the row looks like column labels: all cells non-empty and none
    match the numeric pattern (labels, not data values)."""
    cells = [str(c).strip() for c in row]
    if not cells or any(c == "" for c in cells):
        return False
    return not any(_NUMERIC_RE.match(c) for c in cells)


# ── frame ───────────────────────────────────────────────────────────────────

def _set_table_frame(tbl) -> None:
    """Set table-level borders (none at table level; per-cell controls) + width."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for tag in ("w:tblStyle", "w:tblBorders", "w:tblCellMar"):
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")


# ── per-row population ──────────────────────────────────────────────────────

def _populate_row(tbl, row_idx: int, source_row, n_cols: int, *,
                  is_header: bool, is_last_body: bool,
                  numeric_cols: set[int]) -> None:
    row = tbl.rows[row_idx]
    for col_idx in range(n_cols):
        cell = row.cells[col_idx]
        text = source_row[col_idx] if col_idx < len(source_row) else ""
        _style_cell(cell, text, col_idx=col_idx, row_idx=row_idx,
                    is_header=is_header, is_last_body=is_last_body,
                    is_numeric_col=col_idx in numeric_cols)


def _style_cell(cell, text: str, *, col_idx: int, row_idx: int,
                is_header: bool, is_last_body: bool,
                is_numeric_col: bool) -> None:
    # Only the first header row gets the red Classic fill; subsequent header
    # rows (row_idx > 0 and is_header) are treated as bold body rows so that
    # data rows accidentally placed in table.headers don't render in red.
    true_header = is_header and row_idx == 0

    # Fill — always classic style (red header, zebra body)
    if true_header:
        set_cell_color(cell, T.RED)
    elif row_idx % 2 == 0:
        set_cell_color(cell, T.BG_SUBTLE)
    else:
        set_cell_color(cell, T.WHITE)

    # Borders
    set_cell_borders(
        cell,
        top=T.BORDER_STR if true_header else (T.BORDER_DEF if row_idx > 0 else None),
        bottom=T.BORDER_DEF if is_last_body else None,
        left=None, right=None,
    )

    set_cell_padding(cell, top=T.CELL_PAD_V_TWIPS, bottom=T.CELL_PAD_V_TWIPS,
                     left=T.CELL_PAD_H_TWIPS, right=T.CELL_PAD_H_TWIPS)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Remove default empty para, add our own
    for p in list(cell.paragraphs):
        cell._tc.remove(p._p)
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before_twips=0, after_twips=0,
                          line_multiple=T.TEXT_TABLE.line_spacing)

    # Alignment
    if is_numeric_col and not is_header:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif true_header:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
    elif col_idx == 0:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Badge cell: short status text → bold with sentiment color
    if not is_header and text.strip().lower() in _BADGE_ALL:
        _style_badge_run(para.add_run(text), text.strip().lower())
        return

    run = para.add_run(text)

    if true_header:
        apply_text_style(run, T.TITLES_TABLE)
    elif is_header:
        # Non-first header rows: bold body text (data rows mis-placed in headers)
        apply_text_style(run, T.TEXT_TABLE, override_bold=True)
    elif col_idx == 0:
        apply_text_style(run, T.TEXT_TABLE, override_color=T.TEXT_PRI)
    else:
        apply_text_style(run, T.TEXT_TABLE, override_color=T.TEXT_SEC)


# ── merges ──────────────────────────────────────────────────────────────────

def _apply_merge(tbl, m: MergeSpec, *, n_cols: int, total_rows: int) -> None:
    """Apply a rowspan/colspan to the (m.row, m.col) anchor cell."""
    if m.row >= total_rows or m.col >= n_cols:
        return

    # Column span
    if m.colspan > 1:
        anchor = tbl.rows[m.row].cells[m.col]
        right_bound = min(m.col + m.colspan - 1, n_cols - 1)
        # python-docx Cell.merge will combine — avoids manual gridSpan math.
        right_cell = tbl.rows[m.row].cells[right_bound]
        try:
            anchor.merge(right_cell)
        except Exception:
            # fall back to gridSpan write-through
            set_cell_grid_span(anchor, m.colspan)

    # Row span
    if m.rowspan > 1:
        anchor = tbl.rows[m.row].cells[m.col]
        set_cell_vertical_merge(anchor, restart=True)
        for r in range(m.row + 1, min(m.row + m.rowspan, total_rows)):
            try:
                set_cell_vertical_merge(tbl.rows[r].cells[m.col], restart=False)
            except IndexError:
                break


# ── helpers ─────────────────────────────────────────────────────────────────

def _infer_col_count(table: Table) -> int:
    candidates = []
    for hr in table.headers:
        candidates.append(len(hr))
    for r in table.rows:
        candidates.append(len(r))
    return max(candidates) if candidates else 0


def _numeric_columns(rows, n_cols: int) -> set[int]:
    cols: set[int] = set()
    if not rows:
        return cols
    for c in range(n_cols):
        total = numeric = 0
        for row in rows:
            if c >= len(row):
                continue
            t = (row[c] or "").strip()
            if not t:
                continue
            total += 1
            if _NUMERIC_RE.match(t):
                numeric += 1
        if total > 0 and numeric / total >= 0.5:
            cols.add(c)
    return cols


def _style_badge_run(run, lower_text: str) -> None:
    if lower_text in _BADGE_POSITIVE:
        color = T.TEXT_PRI
    elif lower_text in _BADGE_NEGATIVE:
        color = T.RED
    else:
        color = T.TEXT_SEC
    apply_text_style(run, T.TEXT_TABLE, override_bold=True, override_color=color)


def _add_caption(doc_docx, caption: str) -> None:
    para = doc_docx.add_paragraph()
    set_paragraph_spacing(para, before_twips=80, after_twips=240, line_multiple=1.2)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(caption)
    apply_text_style(run, T.TEXT_DISCLAIMER, override_italic=True)
