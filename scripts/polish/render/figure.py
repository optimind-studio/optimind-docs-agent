"""Figure renderer — inline image with optional caption."""
from __future__ import annotations

import io

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from ..model import Figure
from . import tokens as T
from .xml_utils import apply_text_style, set_paragraph_spacing


MAX_FIGURE_WIDTH_IN = 6.0   # fits within 8.5" page with 1" margins each side


def render(doc_docx, figure: Figure) -> None:
    # Image paragraph
    para = doc_docx.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before_twips=240, after_twips=80, line_multiple=1.2)

    run = para.add_run()
    try:
        run.add_picture(io.BytesIO(figure.image_bytes), width=Inches(MAX_FIGURE_WIDTH_IN))
    except Exception:
        # Fallback: skip image, leave a placeholder note so content-preservation can still
        # reference the surrounding context.
        placeholder = run.add_text("[image]") if hasattr(run, "add_text") else None
        if placeholder is None:
            run.text = "[image]"
        apply_text_style(run, T.TEXT_DISCLAIMER)

    # Caption
    if figure.caption:
        cap = doc_docx.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, after_twips=240, line_multiple=1.2)
        cap_run = cap.add_run(figure.caption)
        apply_text_style(cap_run, T.TEXT_DISCLAIMER, override_italic=True)
