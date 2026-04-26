"""Callout renderer — 4 variants (insight, next_steps, warning, note).

A callout is a borderless 1×1 table with a fill, hosting:
  - First paragraph: UPPERCASE label in the variant's label color.
  - Remaining paragraphs: body in the variant's body color.

Fixed full-width; no outer border; inner padding for breathing room.
Port of the pattern from scripts/polish_doc.py:678-795.
"""
from __future__ import annotations

from docx.enum.table import WD_ALIGN_VERTICAL

from ..model import Callout
from . import tokens as T
from .xml_utils import (
    apply_text_style, set_cell_borders, set_cell_color,
    set_cell_padding, set_paragraph_spacing,
)


def render(doc_docx, callout: Callout) -> None:
    palette = T.CALLOUT_PALETTES[callout.variant]

    # Spacer paragraph above for breathing room.
    spacer = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer, after_twips=0, line_multiple=0.5)

    table = doc_docx.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]

    set_cell_color(cell, palette.fill)
    set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
    set_cell_padding(cell,
                     top=palette.pad_v_twips, bottom=palette.pad_v_twips,
                     left=palette.pad_h_twips, right=palette.pad_h_twips)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Remove the default empty paragraph inside the cell.
    default_p = cell.paragraphs[0]
    cell._tc.remove(default_p._p)

    # Label paragraph
    label_para = cell.add_paragraph()
    set_paragraph_spacing(label_para, after_twips=60, line_multiple=1.2)
    label_run = label_para.add_run(callout.label)
    apply_text_style(label_run, T.LABELS_MAIN, override_color=palette.label_color)

    # Body paragraphs
    for body_para in callout.body:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after_twips=60, line_multiple=T.TEXT_MAIN.line_spacing)
        for r in body_para.runs:
            if not r.text:
                continue
            run = p.add_run(r.text)
            apply_text_style(
                run, T.TEXT_MAIN,
                override_bold=r.bold,
                override_italic=r.italic,
                override_color=palette.body_color,
            )

    # Spacer paragraph below
    spacer_after = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer_after, after_twips=0, line_multiple=0.5)
