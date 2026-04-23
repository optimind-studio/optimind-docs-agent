"""KPI strip renderer — horizontal row of 2–5 branded tiles.

Rendered as a borderless 1-row table. Each cell hosts three stacked paragraphs:
  - Label (Labels/Main uppercase, text/secondary, tracking 1.2px)
  - Value (large, bold, brand red for emphasis)
  - Delta (small, text/secondary) — optional

Why a table and not floating text boxes: tables never overlap other content.
This is exactly why the old screenshots broke — the source had floaters.
"""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ..model import KPIStrip
from . import tokens as T
from .xml_utils import (
    apply_text_style, set_cell_borders, set_cell_color,
    set_cell_padding, set_paragraph_spacing,
)


KPI_VALUE_STYLE = T.TextStyle(
    size=Pt(22), bold=True, color=T.RED, line_spacing=1.1,
)


def render(doc_docx, strip: KPIStrip) -> None:
    if not strip.cards:
        return

    # spacer above
    spacer = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer, after_twips=60, line_multiple=0.5)

    n = len(strip.cards)
    table = doc_docx.add_table(rows=1, cols=n)
    table.autofit = False

    # Borderless, full-width
    _set_frame(table)

    for idx, card in enumerate(strip.cards):
        cell = table.rows[0].cells[idx]
        set_cell_color(cell, T.BG_SUBTLE if idx % 2 == 1 else T.WHITE)
        set_cell_borders(cell, top=T.BORDER_DEF, bottom=T.BORDER_DEF,
                         left=None, right=None)
        set_cell_padding(cell, top=220, bottom=220, left=200, right=200)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Remove default empty paragraph
        for p in list(cell.paragraphs):
            cell._tc.remove(p._p)

        # Label (uppercase)
        label_p = cell.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(label_p, after_twips=40, line_multiple=1.15)
        label_run = label_p.add_run(card.label)
        apply_text_style(label_run, T.LABELS_MAIN)

        # Value
        value_p = cell.add_paragraph()
        value_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(value_p, after_twips=40, line_multiple=1.1)
        val_run = value_p.add_run(card.value)
        apply_text_style(val_run, KPI_VALUE_STYLE)

        # Delta (optional)
        if card.delta:
            delta_p = cell.add_paragraph()
            delta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(delta_p, after_twips=0, line_multiple=1.2)
            delta_run = delta_p.add_run(card.delta)
            apply_text_style(delta_run, T.TEXT_DISCLAIMER)

    # spacer below
    spacer_after = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer_after, after_twips=60, line_multiple=0.5)


def _set_frame(tbl) -> None:
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for tag in ("w:tblStyle", "w:tblBorders", "w:tblCellMar"):
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
