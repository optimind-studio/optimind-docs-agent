"""Top-level renderer — walks canonical Document and writes branded .docx.

Flow:
  1. Render cover (Jinja template → tmp .docx).
  2. Create body doc; set margins; add header/footer.
  3. For each Block in the Document, dispatch to the per-kind renderer.
  4. Merge cover + body + save to output_path.

Table / KPI / Chart renderers are wired up in later steps; for now the writer
routes heading, paragraph, list, callout, figure end-to-end.
"""
from __future__ import annotations

import copy
import os
from pathlib import Path

from docx import Document as _NewDocx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ..model import Document as CanonicalDoc
from . import (
    action_card as action_card_renderer,
    callout as callout_renderer,
    chart as chart_renderer,
    comparison_panel as comparison_panel_renderer,
    cover as cover_mod,
    figure as figure_renderer,
    heading as heading_renderer,
    kpi_strip as kpi_renderer,
    list_block as list_renderer,
    paragraph as paragraph_renderer,
    section_label as section_label_renderer,
    table as table_renderer,
    tokens as T,
)
from .dynamic_dispatch import get_dynamic_renderer
from .xml_utils import apply_text_style, set_paragraph_spacing


STATIC_RENDERERS = {
    "heading", "paragraph", "list", "callout",
    "table", "kpi_strip", "chart", "figure",
    "section_label", "action_card", "comparison_panel",
}


CONFIDENTIALITY = (
    "This document is confidential and intended solely for the addressed recipient. "
    "Do not distribute without prior authorisation."
)


def write(doc: CanonicalDoc, output_path: Path) -> Path:
    """Render `doc` to `output_path` (.docx)."""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    body = _NewDocx()
    _configure_page(body)

    # Walk blocks and dispatch.
    for block in doc.blocks:
        _render_block(body, block, doc.warnings)

    _add_header_footer(body, title=doc.title)

    # Merge cover in front of body.
    cover_path = cover_mod.render_cover(doc.title, doc.client, doc.period)
    try:
        final = _merge_cover_and_body(cover_path, body)
        final.save(str(output_path))
    finally:
        try:
            os.unlink(cover_path)
        except OSError:
            pass

    return output_path


# ── dispatch ────────────────────────────────────────────────────────────────

def _render_block(body_docx, block, warnings: list | None = None) -> None:
    kind = block.kind
    c = block.content
    if kind == "figure":
        # Images are not rendered in v0.5 — skipped cleanly, logged in HTML report.
        if warnings is not None:
            warnings.append(
                f"block {block.source_index}: figure omitted "
                f"(image rendering disabled in v0.5)"
            )
        return
    elif kind == "chart":
        # Charts are not rendered in v0.5 — skipped cleanly, logged in HTML report.
        if warnings is not None:
            warnings.append(
                f"block {block.source_index}: chart omitted "
                f"(chart rendering disabled in v0.5)"
            )
        return
    elif kind == "heading":
        heading_renderer.render(body_docx, c)
    elif kind == "paragraph":
        paragraph_renderer.render(body_docx, c)
    elif kind == "list":
        list_renderer.render(body_docx, c)
    elif kind == "callout":
        callout_renderer.render(body_docx, c)
    elif kind == "table":
        table_renderer.render(body_docx, c)
    elif kind == "kpi_strip":
        kpi_renderer.render(body_docx, c)
    elif kind == "section_label":
        section_label_renderer.render(body_docx, c)
    elif kind == "action_card":
        action_card_renderer.render(body_docx, c)
    elif kind == "comparison_panel":
        comparison_panel_renderer.render(body_docx, c)
    else:
        # Dynamic dispatch: DS-Extender may have generated a renderer for this
        # kind in render/dynamic/. Fall back to a paragraph if unavailable so
        # the pipeline never hard-fails on an unknown kind.
        renderer = get_dynamic_renderer(kind)
        if renderer is not None:
            renderer(body_docx, c)
        else:
            paragraph_renderer.render(body_docx, c)


# ── page setup ──────────────────────────────────────────────────────────────

def _configure_page(body_docx) -> None:
    """Set page margins and default section properties."""
    for section in body_docx.sections:
        section.top_margin    = _twips_to_emu(T.PAGE_MARGIN_TOP_TWIPS)
        section.bottom_margin = _twips_to_emu(T.PAGE_MARGIN_BOTTOM_TWIPS)
        section.left_margin   = _twips_to_emu(T.PAGE_MARGIN_LEFT_TWIPS)
        section.right_margin  = _twips_to_emu(T.PAGE_MARGIN_RIGHT_TWIPS)


def _twips_to_emu(twips: int):
    from docx.shared import Emu
    # 1 twip = 635 EMU (1 inch = 914400 EMU = 1440 twips)
    return Emu(int(twips * 635))


# ── header / footer ─────────────────────────────────────────────────────────

def _add_header_footer(body_docx, title: str) -> None:
    """Add a simple header (title left, page number right) and footer (disclaimer)."""
    for section in body_docx.sections:
        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(hp, after_twips=0, line_multiple=1.0)
        title_run = hp.add_run(title)
        apply_text_style(title_run, T.LABELS_MAIN)

        # Tab stop to push page number to right.
        hp.add_run("\t").bold = False
        page_run = hp.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        page_run._r.append(fld_begin)
        page_run._r.append(instr)
        page_run._r.append(fld_sep)
        page_run._r.append(fld_end)
        apply_text_style(page_run, T.TEXT_DISCLAIMER)

        # Right-aligned tab stop at page inner width
        from docx.shared import Cm
        tab_stops = hp.paragraph_format.tab_stops
        try:
            tab_stops.clear_all()
        except Exception:
            pass
        try:
            tab_stops.add_tab_stop(Cm(16.3), alignment=3)  # 3 = right
        except Exception:
            pass

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(fp, after_twips=0, line_multiple=1.0)
        fr = fp.add_run(CONFIDENTIALITY)
        apply_text_style(fr, T.TEXT_DISCLAIMER)


# ── cover merge ─────────────────────────────────────────────────────────────

def _merge_cover_and_body(cover_path: Path, body_doc) -> object:
    """Prepend the rendered cover to body_doc and return a merged Document.

    Ported from scripts/polish_doc.py:merge_cover_with_body.
    """
    final = _NewDocx(str(cover_path))
    final_body = final.element.body

    # Page break at end of cover body, before its sectPr.
    pg = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    pg.append(r)

    sectPr = final_body.find(qn("w:sectPr"))
    if sectPr is not None:
        final_body.insert(list(final_body).index(sectPr), pg)
    else:
        final_body.append(pg)

    # Copy body content (minus its sectPr).
    src_body = body_doc.element.body
    for elem in list(src_body):
        if elem.tag == qn("w:sectPr"):
            continue
        final_body.append(copy.deepcopy(elem))

    return final
