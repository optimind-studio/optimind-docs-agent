"""Comparison-panel renderer — 2-column What Worked / What Needs Improvement."""
from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from ..model import ComparisonPanel
from . import tokens as T
from .xml_utils import (
    apply_text_style, set_cell_borders, set_cell_color, set_cell_padding,
    set_paragraph_spacing,
)


def render(doc_docx, panel: ComparisonPanel) -> None:
    # Spacer before
    spacer = doc_docx.add_paragraph()
    set_paragraph_spacing(spacer, before_twips=0, after_twips=80, line_multiple=1.0)

    tbl = doc_docx.add_table(rows=1, cols=2)
    tbl.autofit = False
    _set_table_frame(tbl)

    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    # Explicit 50/50 split — prevents Word from auto-collapsing columns
    from .xml_utils import set_cell_width
    set_cell_width(left_cell, 2500, "pct")
    set_cell_width(right_cell, 2500, "pct")

    _fill_column(left_cell, panel.left_title, panel.left_items,
                 bg=T.BG_BRAND, label_color=T.RED)
    _fill_column(right_cell, panel.right_title, panel.right_items,
                 bg=T.BG_SUBTLE, label_color=T.TEXT_PRI)

    # Spacer after
    post = doc_docx.add_paragraph()
    set_paragraph_spacing(post, before_twips=0, after_twips=160, line_multiple=1.0)


def _set_table_frame(tbl) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    for tag in ("w:tblStyle", "w:tblBorders", "w:tblCellMar"):
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")


def _fill_column(cell, title: str, items: list[str], *, bg, label_color) -> None:
    from docx.shared import Pt
    set_cell_color(cell, bg)
    set_cell_borders(cell, top=None, bottom=None, left=None, right=None)
    pad = int(16 * 20)  # 16pt in twips — matches Figma padding
    set_cell_padding(cell, top=pad, bottom=pad, left=pad, right=pad)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Remove default empty para
    for p in list(cell.paragraphs):
        cell._tc.remove(p._p)

    # Label
    label_para = cell.add_paragraph()
    set_paragraph_spacing(label_para, before_twips=0, after_twips=120, line_multiple=1.0)
    label_run = label_para.add_run(title)
    apply_text_style(label_run, T.LABELS_MAIN, override_color=label_color)

    # Bullet items
    for item in items:
        item_para = cell.add_paragraph()
        set_paragraph_spacing(item_para, before_twips=0, after_twips=60, line_multiple=1.3)
        item_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = item_para.add_run(f"• {item}")
        apply_text_style(run, T.TEXT_MAIN)
