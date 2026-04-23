"""
Optimind Document Polisher
Applies Optimind brand styling to a Word document without changing content.

Usage:
  python polish_doc.py --input path/to/file.docx \
                       --title "Document Title" \
                       --client "Client Name" \
                       --period "1 Feb – 28 Feb, 2026"
"""
import argparse
import copy
import json
import os
import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Twips
from docxtpl import DocxTemplate
from lxml import etree

# ── Brand tokens ──────────────────────────────────────────────────────────────
RED          = RGBColor(0xF5, 0x2C, 0x39)
TEXT_PRI     = RGBColor(0x00, 0x00, 0x00)
TEXT_SEC     = RGBColor(0x62, 0x65, 0x67)
BG_SUBTLE    = RGBColor(0xF2, 0xF3, 0xF4)
BG_BRAND     = RGBColor(0xFE, 0xEC, 0xEE)
BORDER_DEF   = RGBColor(0xE5, 0xE7, 0xE9)
BORDER_STR   = RGBColor(0xD7, 0xDB, 0xDD)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

ASSETS = Path(__file__).parent.parent / "assets"
COVER_TEMPLATE = ASSETS / "cover_template.docx"
# Output lives in the user's home drop-folder, not inside the plugin directory.
# Override with OPTIMIND_DOCS_OUTPUT if colleagues need a custom location.
OUTPUT_DIR = Path(os.environ.get(
    "OPTIMIND_DOCS_OUTPUT",
    str(Path.home() / "OptimindDocs" / "output"),
))

HEADING_MAP = {
    1: {"size": Pt(16), "bold": True,  "color": TEXT_PRI},
    2: {"size": Pt(13), "bold": True,  "color": TEXT_PRI},
    3: {"size": Pt(11), "bold": True,  "color": TEXT_PRI},
    4: {"size": Pt(10), "bold": True,  "color": TEXT_PRI},
}

# Detect headings by numbered prefix (e.g. "1. Title", "3.1 Title", "3.1.2 Title")
_HEADING_RE = [
    (re.compile(r'^\d+\.\d+\.\d+[\s\.]'), 3),
    (re.compile(r'^\d+\.\d+[\s\.]'), 2),
    (re.compile(r'^\d+[\.\s]\s*\S'), 1),
]


def _classify_heading(para) -> tuple[str | None, int | None]:
    """Return (source, level) for a heading, or (None, None) if not a heading.
    Source is one of: 'word_style', 'numeric', 'bold_fallback'."""
    text = para.text.strip()
    if not text:
        return None, None

    # Word style takes priority
    style_name = para.style.name if para.style is not None else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
            if level in (1, 2, 3, 4):
                return "word_style", level
        except ValueError:
            pass

    # Numeric prefix
    for pattern, level in _HEADING_RE:
        if pattern.match(text):
            return "numeric", level

    # Short bold paragraph with no numbering → H3 sub-label (ambiguous fallback)
    has_bold = any(r.bold for r in para.runs if r.text.strip())
    if has_bold and len(text) < 60 and '\n' not in text:
        pPr = para._p.find(qn('w:pPr'))
        has_num = pPr is not None and pPr.find(qn('w:numPr')) is not None
        if not has_num:
            return "bold_fallback", 3

    return None, None


def _infer_heading_level(para) -> int | None:
    """Return heading level 1-3 if paragraph looks like a heading, else None."""
    _, level = _classify_heading(para)
    return level

CONFIDENTIALITY = (
    "This document is confidential and intended solely for the addressed recipient. "
    "Do not distribute without prior authorisation."
)


# ── XML helpers ───────────────────────────────────────────────────────────────

def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_color(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color(rgb))
    # Remove existing shd
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    tcPr.append(shd)


def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Pass RGBColor or None to skip."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    sides = {"top": top, "bottom": bottom, "left": left, "right": right}
    for side, color in sides.items():
        el = tcBorders.find(qn(f"w:{side}"))
        if el is not None:
            tcBorders.remove(el)
        if color is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")    # 0.5pt
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), hex_color(color))
            tcBorders.append(el)
        else:
            # Explicitly nil the border so Word hides it
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)


def set_cell_padding(cell, top_twips=100, bottom_twips=100, left_twips=120, right_twips=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top_twips), ("bottom", bottom_twips),
                      ("left", left_twips), ("right", right_twips)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_run_font(run, size: Pt, bold: bool, color: RGBColor, italic=False,
                 letter_spacing_pt: float = 0):
    run.font.name = "Poppins"
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    for tag in [qn("w:rFonts")]:
        el = rPr.find(tag)
        if el is None:
            el = OxmlElement("w:rFonts")
            rPr.insert(0, el)
        el.set(qn("w:ascii"), "Poppins")
        el.set(qn("w:hAnsi"), "Poppins")
        el.set(qn("w:cs"), "Poppins")
    if letter_spacing_pt:
        spacing = rPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            rPr.append(spacing)
        spacing.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
    # Remove RTL marker — prevents number/punctuation reordering in mixed-direction docs
    for rtl_el in rPr.findall(qn("w:rtl")):
        rPr.remove(rtl_el)


def set_run_caps(run, on: bool = True):
    """Display-only uppercase via w:caps — preserves underlying text content."""
    rPr = run._r.get_or_add_rPr()
    for el in rPr.findall(qn("w:caps")):
        rPr.remove(el)
    if on:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "1")
        rPr.append(caps)


def set_para_line_height(para, multiple=1.5):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(int(multiple * 240)))
    spacing.set(qn("w:lineRule"), "auto")


def add_field(paragraph, field_code: str):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar_begin)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_code} "
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run3._r.append(fldChar_end)


# ── Cover handling ─────────────────────────────────────────────────────────────

def render_cover(title: str, client: str, period: str) -> Path:
    """Render the cover template with Jinja vars → temp file path."""
    if not COVER_TEMPLATE.exists():
        raise FileNotFoundError(
            f"Cover template not found at {COVER_TEMPLATE}\n"
            "Run: python scripts/build_cover_template.py"
        )
    tpl = DocxTemplate(str(COVER_TEMPLATE))
    tpl.render({"TITLE": title, "CLIENT": client, "PERIOD": period})
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tpl.save(tmp.name)
    return Path(tmp.name)


def is_cover_page(doc: Document) -> bool:
    """Heuristic: first page is a cover if it has fewer than 6 non-empty paragraphs and no tables."""
    body_paras = [p for p in doc.paragraphs if p.text.strip()]
    # Check for page break within first few paragraphs
    for i, para in enumerate(doc.paragraphs[:10]):
        for run in para.runs:
            if run.break_type is not None:
                return True
        xml = para._p.xml
        if "w:lastRenderedPageBreak" in xml or 'w:type="page"' in xml:
            return True
    # Fallback: very few paragraphs and no tables on first logical page
    return len(body_paras) < 6 and len(doc.tables) == 0


# ── Heading & body styling ────────────────────────────────────────────────────

def _force_ltr(para) -> None:
    """Remove bidi/RTL markers from paragraph and all its runs."""
    pPr = para._p.get_or_add_pPr()
    for el in pPr.findall(qn("w:bidi")):
        pPr.remove(el)
    bidi_off = OxmlElement("w:bidi")
    bidi_off.set(qn("w:val"), "0")
    pPr.append(bidi_off)
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None:
            for el in rPr.findall(qn("w:rtl")):
                rPr.remove(el)


def _get_list_level(para) -> int:
    """Return the list indent level (0-based) for a list paragraph, or 0."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return 0
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return 0
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        return 0
    try:
        return int(ilvl.get(qn("w:val"), 0))
    except (ValueError, TypeError):
        return 0


def _has_numPr(para) -> bool:
    pPr = para._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def style_paragraph(para, doc_title: str = "") -> tuple[bool, str | None]:
    """Apply brand styling to a paragraph. Returns (was_styled, heading_source_or_None)."""
    # Empty paragraphs: collapse to zero height so they don't create phantom dividers
    if not para.text.strip():
        pPr = para._p.get_or_add_pPr()
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr.append(sp)
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        sp.set(qn("w:line"), str(int(1.0 * 240)))
        sp.set(qn("w:lineRule"), "exact")
        return False, None

    style_name = para.style.name

    heading_source, level = _classify_heading(para)

    if level and level in HEADING_MAP:
        cfg = HEADING_MAP[level]
        for run in para.runs:
            set_run_font(run, cfg["size"], cfg["bold"], cfg["color"])
        set_para_line_height(para)
        # Force LTR so number prefixes (e.g. "1. ", "3.1 ") don't get reordered
        pPr = para._p.get_or_add_pPr()
        for bidi_el in pPr.findall(qn("w:bidi")):
            pPr.remove(bidi_el)
        bidi_off = OxmlElement("w:bidi")
        bidi_off.set(qn("w:val"), "0")
        pPr.append(bidi_off)
        # Remove any existing pBdr (e.g. VML rule remnants or prior runs)
        for existing_bdr in pPr.findall(qn("w:pBdr")):
            pPr.remove(existing_bdr)
        if level == 1:
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after  = Pt(8)
            para.paragraph_format.left_indent  = Pt(0)
        elif level == 2:
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after  = Pt(4)
            para.paragraph_format.left_indent  = Pt(0)
        else:
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after  = Pt(2)
            para.paragraph_format.left_indent  = Pt(0)
        return True, heading_source

    # Disclaimer / caption
    if any(style_name.lower().startswith(s) for s in ["caption", "footnote", "endnote"]):
        for run in para.runs:
            set_run_font(run, Pt(9), False, TEXT_SEC)
        set_para_line_height(para)
        return True, None

    # List / bullet paragraphs
    if _has_numPr(para) or "list" in style_name.lower():
        ilvl = _get_list_level(para)
        for run in para.runs:
            set_run_font(run, Pt(11), False, TEXT_SEC)
        set_para_line_height(para)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(3)
        # Normalise indent: override the numbering definition's large defaults
        # left=360+ilvl*360 twips, hanging=180 twips — tight, consistent indent
        pPr = para._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        left_twips = str(360 + ilvl * 360)
        ind.set(qn("w:left"),    left_twips)
        ind.set(qn("w:hanging"), "180")
        _force_ltr(para)
        return True, None

    # Normal body text
    if style_name in ("Normal", "Body Text", "Default Paragraph Style", "") or \
       style_name.lower().startswith("body"):
        for run in para.runs:
            set_run_font(run, Pt(11), False, TEXT_SEC)
        set_para_line_height(para)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(4)
        _force_ltr(para)
        return True, None

    return False, None


def count_styled_headings(doc: Document) -> int:
    count = 0
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") or _infer_heading_level(para) is not None:
            count += 1
    return count


# ── Table styling ─────────────────────────────────────────────────────────────

_NUMERIC_RE = re.compile(r'^\s*[\$€£¥]?\s*[-+]?\d[\d,]*(\.\d+)?\s*%?\s*$')


def _is_numeric_cell(text: str) -> bool:
    return bool(_NUMERIC_RE.match(text.strip()))


def _table_is_numeric(table) -> bool:
    """Heuristic: ≥60% of non-header cells parse as numbers / percentages / currency."""
    if len(table.rows) < 2:
        return False
    total = 0
    numeric = 0
    for row in table.rows[1:]:  # skip header row
        for cell in row.cells:
            text = cell.text
            if not text.strip():
                continue
            total += 1
            if _is_numeric_cell(text):
                numeric += 1
    return total > 0 and (numeric / total) >= 0.6


def style_table_minimal(table) -> None:
    """Minimal table variant: no fills, top/header-bottom rules (strong),
    row dividers (default). Numeric columns right-aligned."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    for tag in [qn("w:tblStyle"), qn("w:tblBorders"), qn("w:tblCellMar")]:
        el = tblPr.find(tag)
        if el is not None:
            tblPr.remove(el)

    # No table-level borders (we apply per-cell borders for fine control)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    last_row_idx = len(table.rows) - 1

    # Determine numeric columns (column-wise: majority numeric)
    col_count = len(table.rows[0].cells) if table.rows else 0
    numeric_cols = set()
    for col_idx in range(col_count):
        total = 0
        numeric = 0
        for row in table.rows[1:]:
            if col_idx >= len(row.cells):
                continue
            t = row.cells[col_idx].text
            if not t.strip():
                continue
            total += 1
            if _is_numeric_cell(t):
                numeric += 1
        if total > 0 and (numeric / total) >= 0.5:
            numeric_cols.add(col_idx)

    for row_idx, row in enumerate(table.rows):
        is_header = row_idx == 0
        is_last   = row_idx == last_row_idx

        for col_idx, cell in enumerate(row.cells):
            # No fill
            set_cell_color(cell, WHITE)

            # Borders: top on header, bottom on header, row dividers, bottom on last row
            top = BORDER_STR if is_header else None
            if is_header:
                bottom = BORDER_STR
            elif is_last:
                bottom = BORDER_STR
            else:
                bottom = BORDER_DEF
            set_cell_borders(cell, top=top, bottom=bottom, left=None, right=None)

            set_cell_padding(cell, top_twips=100, bottom_twips=100,
                             left_twips=120, right_twips=120)

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            is_numeric = col_idx in numeric_cols
            for para in cell.paragraphs:
                if is_numeric:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    if is_header:
                        set_run_font(run, Pt(10), True, TEXT_SEC, letter_spacing_pt=1.2)
                        set_run_caps(run, True)
                    elif col_idx == 0:
                        set_run_font(run, Pt(10), False, TEXT_PRI)
                    else:
                        set_run_font(run, Pt(10), False, TEXT_SEC)
                set_para_line_height(para, multiple=1.25)


def style_table(table) -> None:
    """Apply Optimind branded table style: red header, alternating rows."""
    # Remove existing table style
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Remove existing tblStyle and tblBorders
    for tag in [qn("w:tblStyle"), qn("w:tblBorders"), qn("w:tblCellMar")]:
        el = tblPr.find(tag)
        if el is not None:
            tblPr.remove(el)

    # Table-level: no outside borders
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Table width: 100%
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    last_row_idx = len(table.rows) - 1

    for row_idx, row in enumerate(table.rows):
        is_header = row_idx == 0
        is_last   = row_idx == last_row_idx
        is_alt = (row_idx % 2 == 0) and not is_header

        for col_idx, cell in enumerate(row.cells):
            # Background fill
            if is_header:
                set_cell_color(cell, RED)
            elif is_alt:
                set_cell_color(cell, BG_SUBTLE)
            else:
                set_cell_color(cell, WHITE)

            # Borders
            set_cell_borders(
                cell,
                top=BORDER_STR if is_header else (BORDER_DEF if row_idx > 0 else None),
                bottom=BORDER_DEF if is_last else None,
                left=None,
                right=None,
            )

            # Padding
            set_cell_padding(cell, top_twips=100, bottom_twips=100,
                             left_twips=120, right_twips=120)

            # Vertical alignment: center
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Text styling
            for para in cell.paragraphs:
                align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                para.alignment = align
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    if is_header:
                        set_run_font(run, Pt(10), True, WHITE)
                    elif col_idx == 0:
                        set_run_font(run, Pt(10), False, TEXT_PRI)
                    else:
                        set_run_font(run, Pt(10), False, TEXT_SEC)
                set_para_line_height(para, multiple=1.25)


# ── Header & footer ───────────────────────────────────────────────────────────

def add_header_footer(doc: Document, doc_title: str) -> None:
    """Add branded header (title + page number) and footer (disclaimer) to all sections."""
    for i, section in enumerate(doc.sections):
        section.different_first_page_header_footer = True

        # ── Header ────────────────────────────────────────────────────────────
        header = section.header
        # Clear existing content
        for para in header.paragraphs:
            for run in para.runs:
                run.text = ""
        if not header.paragraphs:
            header.add_paragraph()

        hpara = header.paragraphs[0]
        hpara.clear()
        hpara.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Title left — Labels/Main: SemiBold 10pt, TEXT_SEC, UPPERCASE, tracking 1.2px
        title_run = hpara.add_run(doc_title.upper())
        set_run_font(title_run, Pt(10), True, TEXT_SEC, letter_spacing_pt=1.2)

        # Tab stop → right-align page number
        pPr = hpara._p.get_or_add_pPr()
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.append(tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9356")  # ~16.5cm
        tabs.append(tab)

        tab_run = hpara.add_run("\t")
        page_run = hpara.add_run("PAGE ")
        set_run_font(page_run, Pt(10), True, TEXT_SEC, letter_spacing_pt=1.2)

        add_field(hpara, "PAGE")

        # No bottom border on header — conflicts with H1 section dividers on pages
        # where an H1 falls near the top, producing two lines instead of one.

        # ── Footer ────────────────────────────────────────────────────────────
        footer = section.footer
        for para in footer.paragraphs:
            for run in para.runs:
                run.text = ""
        if not footer.paragraphs:
            footer.add_paragraph()

        fpara = footer.paragraphs[0]
        fpara.clear()
        fpara.alignment = WD_ALIGN_PARAGRAPH.LEFT

        disc_run = fpara.add_run(CONFIDENTIALITY)
        set_run_font(disc_run, Pt(9), False, TEXT_SEC)
        set_para_line_height(fpara)

        # No top border on footer — collides with H1 section divider on the next page.


# ── Callout detection & styling ───────────────────────────────────────────────

def _get_shading(para) -> str | None:
    """Return fill hex string if paragraph has a shaded background, else None."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    return fill if fill and fill not in ("auto", "FFFFFF", "ffffff", "") else None


def _build_callout_table(bg_color: RGBColor, py_twips: int) -> tuple:
    """Build the XML skeleton for a borderless single-cell callout table.
    Returns (tbl, tc) — caller appends paragraphs to tc."""
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000"); tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")

    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), "5000"); tcW.set(qn("w:type"), "pct")
    tcPr.append(tcW)

    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        el = OxmlElement(f"w:{side}"); el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color(bg_color))
    tcPr.append(shd)

    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", py_twips), ("bottom", py_twips),
                      ("left", 360), ("right", 360)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)
    tc.append(tcPr)

    tr.append(tc); tbl.append(tr)
    return tbl, tc


def _style_callout_paragraph_block(block, is_brand: bool) -> None:
    """Apply label + body styling to a list of paragraphs forming one callout.
    First paragraph's runs become the label (uppercase via w:caps); rest is body."""
    label_color = RED if is_brand else TEXT_PRI
    body_color  = TEXT_PRI if is_brand else TEXT_SEC

    for run in block[0].runs:
        set_run_font(run, Pt(10), True, label_color, letter_spacing_pt=1.2)
        set_run_caps(run, True)
    for para in block[1:]:
        for run in para.runs:
            set_run_font(run, Pt(11), False, body_color)


def _move_para_into_cell(para, tc, body, is_last: bool) -> None:
    p_el = para._p
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr"); p_el.insert(0, pPr)
    for s in pPr.findall(qn("w:shd")):
        pPr.remove(s)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing"); pPr.append(sp)
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0" if is_last else str(int(6 * 20)))
    sp.set(qn("w:line"), str(int(1.5 * 240)))
    sp.set(qn("w:lineRule"), "auto")
    body.remove(p_el)
    tc.append(p_el)


def style_callouts(doc: Document) -> dict:
    """Wrap contiguous shaded paragraph blocks in borderless single-cell tables.
    Returns {'shaded': N}."""
    counts = {"shaded": 0}
    body = doc.element.body
    paras = list(doc.paragraphs)
    i = 0

    while i < len(paras):
        fill = _get_shading(paras[i])
        if fill:
            fill_upper = fill.upper()
            is_brand = fill_upper in ("FEECEE", "FEE", "FFE5E5", "FDDDE0")

            block = []
            j = i
            while j < len(paras) and _get_shading(paras[j]) == fill:
                block.append(paras[j])
                j += 1

            if block:
                counts["shaded"] += 1
                bg_color = BG_BRAND if is_brand else BG_SUBTLE
                py_twips = int(16 * 20) if is_brand else int(14 * 20)

                _style_callout_paragraph_block(block, is_brand)

                insert_idx = list(body).index(block[0]._p)
                tbl, tc = _build_callout_table(bg_color, py_twips)
                for k, p in enumerate(block):
                    _move_para_into_cell(p, tc, body, is_last=(k == len(block) - 1))
                body.insert(insert_idx, tbl)

            i = j
        else:
            i += 1

    return counts


# ── Document merge ────────────────────────────────────────────────────────────

def merge_cover_with_body(cover_path: Path, body_doc: Document) -> Document:
    """Prepend the rendered cover to the body document via XML manipulation."""
    cover_doc = Document(str(cover_path))

    # We'll build the final doc starting from the cover
    final = Document(str(cover_path))
    final_body = final.element.body

    # Insert a page break at the end of cover body
    pg_break_para = OxmlElement("w:p")
    pg_break_r = OxmlElement("w:r")
    pg_break_br = OxmlElement("w:br")
    pg_break_br.set(qn("w:type"), "page")
    pg_break_r.append(pg_break_br)
    pg_break_para.append(pg_break_r)

    # Find sectPr (last element in body) and insert before it
    sectPr = final_body.find(qn("w:sectPr"))
    if sectPr is not None:
        final_body.insert(list(final_body).index(sectPr), pg_break_para)
    else:
        final_body.append(pg_break_para)

    # Copy all body elements from the content doc
    body_doc_body = body_doc.element.body
    content_sectPr = body_doc_body.find(qn("w:sectPr"))

    for elem in list(body_doc_body):
        if elem.tag == qn("w:sectPr"):
            continue  # skip source sectPr
        final_body.append(copy.deepcopy(elem))

    return final


# ── Main ──────────────────────────────────────────────────────────────────────

_FIRST_HEADING_RE = re.compile(r'^\d+[\.\s]\s*\S')


def remove_vml_horizontal_rules(doc: Document) -> int:
    """Remove paragraphs that contain only a VML horizontal rule (v:rect o:hr='t').
    These are decorative dividers baked into the source doc that conflict with our
    own H1 top-border rule.  Returns the number removed."""
    body = doc.element.body
    HR_NS = "urn:schemas-microsoft-com:vml"
    to_remove = []
    for p_el in body.iter(qn("w:p")):
        # A VML hr paragraph has no w:t text but contains a v:rect with o:hr
        has_text = bool("".join(t.text or "" for t in p_el.iter(qn("w:t"))).strip())
        if has_text:
            continue
        for rect in p_el.iter(f"{{{HR_NS}}}rect"):
            o_ns = "urn:schemas-microsoft-com:office:office"
            if rect.get(f"{{{o_ns}}}hr") == "t":
                to_remove.append(p_el)
                break
    for p_el in to_remove:
        parent = p_el.getparent()
        if parent is not None:
            parent.remove(p_el)
    return len(to_remove)


def _para_text(p_el) -> str:
    """Extract plain text from a w:p element."""
    return "".join(
        t.text or "" for t in p_el.iter(qn("w:t"))
    ).strip()


def detect_cover(doc: Document) -> int:
    """Returns the index of the last paragraph that is part of the cover page, or -1.

    Detects cover by (in order of priority):
    1. Explicit page break in the first 15 paragraphs.
    2. The paragraph just before the first numbered H1 heading (e.g. "1. Executive Summary"),
       when that heading appears within the first 30 paragraphs — handles docs with no
       page break between cover metadata and body.
    3. A tiny document with fewer than 6 non-empty paragraphs and no tables.
    """
    paragraphs = list(doc.paragraphs)

    # 1. Explicit page break
    for i, para in enumerate(paragraphs[:15]):
        xml = para._p.xml
        if 'w:type="page"' in xml or "w:pageBreakBefore" in xml:
            return i
        for run in para.runs:
            if 'w:type="page"' in run._r.xml:
                return i

    # 2. First numbered H1 heading — everything before it is cover metadata
    for i, para in enumerate(paragraphs[:30]):
        text = para.text.strip()
        if text and _FIRST_HEADING_RE.match(text):
            if i > 0:
                return i - 1
            return -1

    # 3. Tiny doc fallback
    non_empty = [p for p in paragraphs if p.text.strip()]
    if len(non_empty) < 6 and not doc.tables:
        return len(paragraphs) - 1

    return -1


def strip_cover_elements(doc: Document, cover_end_para_idx: int) -> None:
    """Remove ALL body children (paragraphs AND tables) that appear before or at
    the cover_end_para_idx-th paragraph.  This handles cover tables (e.g. coloured
    callout blocks) that sit between the cover paragraphs."""
    body = doc.element.body
    paragraphs = body.findall(qn("w:p"))
    if cover_end_para_idx < 0 or cover_end_para_idx >= len(paragraphs):
        return
    # Find the w:p element that marks the end of the cover
    last_cover_p = paragraphs[cover_end_para_idx]
    # Collect every body child up to AND including that element
    to_remove = []
    for child in list(body):
        to_remove.append(child)
        if child is last_cover_p:
            break
    for el in to_remove:
        if el.tag != qn("w:sectPr"):   # never remove section properties
            body.remove(el)


def _extract_all_text(source) -> str:
    """Concatenated text of a doc, joined at paragraph / cell boundaries so
    adjacent paragraphs never glue words together. Punctuation and extra
    whitespace collapse to single spaces."""
    if isinstance(source, (str, Path)):
        doc = Document(str(source))
        root = doc.element.body
    else:
        root = source.element.body

    # Walk every w:p element; for each, concatenate its inner w:t runs; join
    # paragraphs with newlines so boundaries survive whitespace collapse.
    paras = []
    for p in root.iter(qn("w:p")):
        chunks = []
        for t in p.iter(qn("w:t")):
            if t.text:
                chunks.append(t.text)
        if chunks:
            paras.append("".join(chunks))

    raw = "\n".join(paras)
    # Collapse any run of non-alphanumeric characters to a single space.
    return re.sub(r"[^A-Za-z0-9]+", " ", raw).strip().lower()


def verify_text_preserved(body_doc: Document, final_doc: Document) -> None:
    """Raise RuntimeError if any body text was dropped during the final
    cover-merge / save step. We compare the *post-cover-strip body* (where all
    our transformations have just finished) against the *final merged doc*.
    Every multi-letter word in the body must appear in the final at least as
    many times as in the body. The cover/header/footer inject extra words —
    that's fine; they only increase output counts."""
    src_n = _extract_all_text(body_doc)
    out_n = _extract_all_text(final_doc)

    if not src_n:
        return

    from collections import Counter
    # Only check words of length ≥ 3 so short fragments (a, of, to, 1, 2…) don't
    # dominate the signal.
    src_words = Counter(w for w in src_n.split() if len(w) >= 3)
    out_words = Counter(w for w in out_n.split() if len(w) >= 3)

    missing = []
    for w, n in src_words.items():
        if out_words[w] < n:
            missing.append((w, n, out_words[w]))

    if missing:
        sample = missing[:5]
        raise RuntimeError(
            f"Content-preservation check failed: {len(missing)} body word(s) "
            f"missing or reduced in final output. Sample: {sample}"
        )


def process(input_path: str, title: str, client: str, period: str,
            table_style: str = "classic") -> dict:
    summary = {
        "title": title, "client": client, "period": period,
        "headings": 0, "tables": 0, "callouts": 0, "output": "",
        "detection": {
            "headings_by_source": {},
            "callouts_by_source": {"shaded": 0, "label_prefix": 0, "single_cell_table": 0},
            "tables_by_variant": {"classic": 0, "minimal": 0},
            "ambiguous_paragraphs": 0,
        },
    }

    doc = Document(input_path)

    # ── Step B: Strip existing cover ─────────────────────────────────────────
    cover_end = detect_cover(doc)
    if cover_end >= 0:
        strip_cover_elements(doc, cover_end)

    # ── Strip VML horizontal rules (original doc dividers we replace with H1 borders)
    remove_vml_horizontal_rules(doc)

    # ── Step C: Apply paragraph styles ───────────────────────────────────────
    heading_count = 0
    headings_by_source: dict[str, int] = {}
    paras = list(doc.paragraphs)
    for para in paras:
        styled, heading_source = style_paragraph(para, title)
        if styled and heading_source is not None:
            heading_count += 1
            headings_by_source[heading_source] = headings_by_source.get(heading_source, 0) + 1
            if heading_source == "bold_fallback":
                summary["detection"]["ambiguous_paragraphs"] += 1
    summary["headings"] = heading_count
    summary["detection"]["headings_by_source"] = headings_by_source

    # ── Step C2: Add divider above each H1 as bottom-border on preceding paragraph
    # Placing the rule on the paragraph BEFORE H1 means it always sits ABOVE the
    # heading regardless of page breaks — no clash with page-boundary lines.
    for i, para in enumerate(paras):
        if _infer_heading_level(para) == 1 and i > 0:
            prev = paras[i - 1]
            pPr = prev._p.get_or_add_pPr()
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                pBdr = OxmlElement("w:pBdr")
                pPr.append(pBdr)
            for existing in pBdr.findall(qn("w:bottom")):
                pBdr.remove(existing)
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:space"), "6")
            bot.set(qn("w:color"), hex_color(BORDER_STR))
            pBdr.append(bot)

    # ── Step F: Callouts (run BEFORE table styling so 1×1 callout tables don't
    #    also get the Classic red-header treatment).
    callout_counts = style_callouts(doc)
    summary["detection"]["callouts_by_source"] = callout_counts
    summary["callouts"] = sum(callout_counts.values())

    # ── Step D: Restyle data tables. Skip 1×1 tables — those are callouts just
    #    created above (or legitimate 1-cell boxes that shouldn't get a red header).
    tables_by_variant = {"classic": 0, "minimal": 0}
    remaining_tables = [
        t for t in doc.tables
        if not (len(t.rows) == 1 and len(t.rows[0].cells) == 1)
    ]
    for table in remaining_tables:
        if table_style == "minimal":
            use_minimal = True
        elif table_style == "auto":
            use_minimal = _table_is_numeric(table)
        else:  # classic
            use_minimal = False

        if use_minimal:
            style_table_minimal(table)
            tables_by_variant["minimal"] += 1
        else:
            style_table(table)
            tables_by_variant["classic"] += 1
    summary["tables"] = len(remaining_tables)
    summary["detection"]["tables_by_variant"] = tables_by_variant

    # ── Step E: Header & footer ───────────────────────────────────────────────
    add_header_footer(doc, title)

    # ── Step A+G: Render cover, merge, save ───────────────────────────────────
    OUTPUT_DIR.mkdir(exist_ok=True, parents=True)
    out_name = Path(input_path).name
    out_path = OUTPUT_DIR / out_name

    cover_path = render_cover(title, client, period)
    try:
        final_doc = merge_cover_with_body(cover_path, doc)

        # ── Content-preservation guardrail: body text must survive the merge.
        verify_text_preserved(doc, final_doc)

        final_doc.save(str(out_path))
    finally:
        os.unlink(cover_path)

    summary["output"] = str(out_path)
    return summary


def main():
    parser = argparse.ArgumentParser(description="Optimind Document Polisher")
    parser.add_argument("--input",  required=True, help="Path to input .docx")
    parser.add_argument("--title",  required=True, help="Document title for cover")
    parser.add_argument("--client", required=True, help="Client name for cover")
    parser.add_argument("--period", required=True, help="Reporting period for cover")
    parser.add_argument("--table-style", choices=["classic", "minimal", "auto"],
                        default="classic",
                        help="Table variant: 'classic' (red header), 'minimal' (rule-based), "
                             "or 'auto' (minimal for numeric tables, classic otherwise).")
    args = parser.parse_args()

    if not Path(args.input).exists():
        print(json.dumps({"error": f"File not found: {args.input}"}))
        sys.exit(1)

    try:
        result = process(args.input, args.title, args.client, args.period,
                         table_style=args.table_style)
    except RuntimeError as e:
        print(json.dumps({"error": str(e)}, ensure_ascii=False))
        sys.exit(2)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
